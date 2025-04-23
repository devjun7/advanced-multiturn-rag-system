"""
Handles loading and processing of text and documents for the RAG system,
including text extraction from files and chunking.
"""
import io
import logging
import os
from typing import List, Optional
import streamlit as st # Added import for st.error/st.warning

# Optional dependencies
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Configure logging
logger = logging.getLogger(__name__)

# --- Configuration for Text Splitting ---
# These values can be tuned based on the embedding model and expected document structure.
# See: https://python.langchain.com/v0.2/docs/how_to/recursive_text_splitter/
from config.settings import CHUNK_SIZE, CHUNK_OVERLAP

# Initialize the text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    length_function=len,
    is_separator_regex=False,
)

def process_text_to_docs(text: str, source: str = "text_input") -> List[Document]:
    """
    Converts raw text into a list of Langchain Document objects using recursive splitting.

    Args:
        text: The input text string.
        source: Identifier for the source of the text (e.g., filename, 'text_input').

    Returns:
        A list of Document objects, where each document is a chunk of the original text.
        Returns an empty list if the input text is empty or None.
    """
    if not text:
        logger.warning("process_text_to_docs called with empty text.")
        return []

    logger.info(f"Splitting text from source '{source}' into chunks...")
    # Use the pre-configured text splitter
    chunks = text_splitter.split_text(text)

    docs = [
        Document(page_content=chunk, metadata={"source": source})
        for chunk in chunks if chunk.strip() # Ensure chunks are not just whitespace
    ]
    logger.info(f"Split text from source '{source}' into {len(docs)} documents (chunks). Chunk size: {CHUNK_SIZE}, Overlap: {CHUNK_OVERLAP}")
    return docs

def process_file_to_docs(uploaded_file) -> List[Document]:
    """
    Extracts text from an uploaded file object (PDF, DOCX, TXT) and processes it into Document chunks.

    Args:
        uploaded_file: A file-like object, typically from Streamlit's file_uploader.
                         Should have 'name' and 'getvalue()' attributes.

    Returns:
        A list of Document objects representing chunks of the extracted text.
        Returns an empty list if the file is None, unsupported, or processing fails.
    """
    if uploaded_file is None:
        logger.warning("process_file_to_docs called with None file.")
        return []

    file_name = uploaded_file.name
    try:
        file_content = uploaded_file.getvalue()
        logger.info(f"Processing file: {file_name}")
    except Exception as e:
        logger.error(f"Could not read content from uploaded file '{file_name}': {e}", exc_info=True)
        st.error(f"Could not read content from file '{file_name}': {e}")
        return []

    docs = []
    extracted_text = ""

    try:
        # PDF Processing
        if file_name.lower().endswith(".pdf"):
            if pypdf is None:
                logger.error("Cannot process PDF '{file_name}', 'pypdf' library is not installed.")
                st.error("The 'pypdf' library is not installed for PDF processing.")
                return []
            extracted_text = _extract_text_from_pdf(file_content, file_name)

        # DOCX Processing
        elif file_name.lower().endswith(".docx"):
            if docx is None:
                logger.error("Cannot process DOCX '{file_name}', 'python-docx' library is not installed.")
                st.error("The 'python-docx' library is not installed for Word document processing.")
                return []
            extracted_text = _extract_text_from_docx(file_content, file_name)

        # TXT Processing (with encoding fallback)
        elif file_name.lower().endswith(".txt"):
            extracted_text = _extract_text_from_txt(file_content, file_name)

        # Attempt to process unknown types as text
        else:
            logger.warning(f"Unsupported file type: {file_name}. Attempting to process as text.")
            extracted_text = _extract_text_from_txt(file_content, file_name, is_fallback=True)

        # If text was successfully extracted, split it into Document chunks
        if extracted_text:
            docs = process_text_to_docs(extracted_text, source=file_name)
        else:
            # Warning/error logged within extraction functions if text is empty
            logger.warning(f"No text extracted from file '{file_name}'. No documents generated.")
            st.warning(f"Could not extract text from file '{file_name}'.")

    except Exception as e:
        logger.error(f"Error processing file '{file_name}': {e}", exc_info=True)
        st.error(f"Error occurred while processing file '{file_name}': {e}")
        return [] # Return empty list on failure

    return docs

# --- Private Helper Functions for Text Extraction ---

def _extract_text_from_pdf(content_bytes: bytes, filename: str) -> str:
    """Extracts text from PDF bytes using pypdf."""
    extracted_text = ""
    try:
        reader = pypdf.PdfReader(io.BytesIO(content_bytes))
        text_list = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text: # Only add if text is extracted
                    text_list.append(page_text)
            except Exception as page_e:
                logger.warning(f"Error extracting text from page {page_num + 1} of PDF '{filename}': {page_e}")
        extracted_text = "\n\n".join(text_list) # Join pages with double newline
        if not extracted_text:
            logger.warning(f"No text could be extracted from PDF '{filename}'. It might be image-based or encrypted.")
    except Exception as e:
        logger.error(f"Failed to read PDF file '{filename}': {e}", exc_info=True)
        st.error(f"Error occurred while reading PDF file '{filename}': {e}")
    return extracted_text

def _extract_text_from_docx(content_bytes: bytes, filename: str) -> str:
    """Extracts text from DOCX bytes using python-docx."""
    extracted_text = ""
    try:
        document = docx.Document(io.BytesIO(content_bytes))
        extracted_text = "\n\n".join([para.text for para in document.paragraphs if para.text.strip()])
        if not extracted_text:
            logger.warning(f"No text could be extracted from DOCX file '{filename}'.")
    except Exception as e:
        logger.error(f"Failed to read DOCX file '{filename}': {e}", exc_info=True)
        st.error(f"Error occurred while reading DOCX file '{filename}': {e}")
    return extracted_text

def _extract_text_from_txt(content_bytes: bytes, filename: str, is_fallback: bool = False) -> str:
    """Extracts text from TXT bytes, attempting UTF-8 and then CP949 decoding."""
    extracted_text = ""
    try:
        extracted_text = content_bytes.decode("utf-8")
        if is_fallback:
            logger.info(f"Successfully processed unknown file type '{filename}' as UTF-8 text.")
    except UnicodeDecodeError:
        logger.warning(f"Could not decode '{filename}' as UTF-8. Trying CP949 (Korean)...") # Keep explanation
        try:
            extracted_text = content_bytes.decode("cp949")
            logger.info(f"Successfully decoded '{filename}' as CP949.")
        except Exception as decode_e:
            logger.error(f"Failed to decode '{filename}' with both UTF-8 and CP949: {decode_e}")
            st.error(f"Failed to decode file '{filename}' with both UTF-8 and CP949: {decode_e}")
            extracted_text = "" # Ensure empty string on failure
    except Exception as e:
       logger.error(f"Error reading text file '{filename}': {e}", exc_info=True)
       extracted_text = ""
    return extracted_text
